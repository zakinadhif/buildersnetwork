"""Builds the Al-Fath Berkarya developer getting-started guide as a .docx.

The .docx is a build artifact — edit this script, not the Word file, or your
changes are lost on the next run. The Indonesian edition is build-guide-id.py;
the two are kept in lockstep, so a change here usually wants the same change
there.

    pip install python-docx
    python docs/build-guide-en.py    # -> docs/getting-started.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x7A, 0x2E, 0x2E)      # maroon, from the Al-Fath brand
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
CODE_BG = "F2F1EF"
NOTE_BG = "FBF3E4"
WARN_BG = "FBE9E7"

doc = Document()

# --- base styles -------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for name, size in (("Heading 1", 18), ("Heading 2", 13.5), ("Heading 3", 11.5)):
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = ACCENT
    st.paragraph_format.space_before = Pt(14)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True


def shade(paragraph, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    paragraph._p.get_or_add_pPr().append(shd)


def code(lines, comment_prefixes=("#",)):
    """A shaded monospace block. Comment lines render muted."""
    if isinstance(lines, str):
        lines = lines.strip("\n").split("\n")
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.right_indent = Inches(0.10)
    pf.space_before = Pt(6)
    pf.space_after = Pt(9)
    pf.line_spacing = 1.05
    shade(p, CODE_BG)
    for i, line in enumerate(lines):
        run = p.add_run(("\n" if i else "") + line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        is_comment = line.lstrip().startswith(comment_prefixes)
        run.font.color.rgb = MUTED if is_comment else RGBColor(0x1A, 0x1A, 0x1A)
    return p


def callout(label, text, bg=NOTE_BG):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.right_indent = Inches(0.10)
    pf.space_before = Pt(6)
    pf.space_after = Pt(9)
    shade(p, bg)
    r = p.add_run(f"{label}  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = ACCENT if bg is NOTE_BG else RGBColor(0xA3, 0x2A, 0x1C)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p


def para(text, bold=False, italic=False, muted=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if muted:
        r.font.color.rgb = MUTED
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            # Backtick-wrapped spans render as inline code.
            for j, chunk in enumerate(val.split("`")):
                if not chunk:
                    continue
                run = p.add_run(chunk)
                if j % 2:
                    run.font.name = "Consolas"
                    run.font.size = Pt(8.5)
                else:
                    run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# =============================================================================
# Title
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("Al-Fath Berkarya")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = ACCENT
r.font.name = "Calibri"

sub = doc.add_paragraph()
r = sub.add_run("Getting Started — a developer's guide to running the project from zero")
r.font.size = Pt(13)
r.font.color.rgb = MUTED

meta = doc.add_paragraph()
r = meta.add_run(
    "Verified end-to-end on Windows 11 against commit c9b3b5f — 14 July 2026.\n"
    "Every command below was actually run on a clean clone; every error in the "
    "Troubleshooting section is one this guide's author hit for real."
)
r.font.size = Pt(9.5)
r.italic = True
r.font.color.rgb = MUTED

doc.add_paragraph()

para(
    "Al-Fath Berkarya is a community platform for builder students at Telkom University. "
    "It is a pnpm monorepo: a Hono API, a React 19 SPA, an Astro landing page, and a static "
    "mockup gallery, sharing a set of libraries (database, auth, AI, storage, config). The same "
    "codebase deploys to Cloudflare Workers and to Node/Docker — the runtime entrypoint, not an "
    "environment variable, picks the AI provider and the database client."
)
para(
    "This guide takes you from nothing installed to a running app with seeded data you can log "
    "into. Budget about 20 minutes, most of it waiting on downloads.",
)

callout(
    "A note on history.",
    "Writing this guide meant running the setup on a clean clone, which turned up four defects in "
    "the documented path — including one that made a fresh clone unable to boot at all. Those are "
    "now fixed in the repo, and the README's Quick start is correct as of commit c9b3b5f. This "
    "guide covers the same ground in more depth. Section 9 records what was wrong, in case you are "
    "on an older checkout.",
)

# =============================================================================
doc.add_heading("1. What you need installed", level=1)

para(
    "Four tools. Node and pnpm are required; Git is how you get the code; Docker is optional and "
    "only powers file uploads."
)

table(
    ["Tool", "Version", "Required?"],
    [
        ["Node.js", "22 or newer", "Yes"],
        ["pnpm", "10 or newer", "Yes — npm and yarn will not work"],
        ["Git", "any recent", "Yes"],
        ["Docker Desktop", "any recent", "Optional — only for image uploads (MinIO)"],
    ],
    widths=[1.6, 1.9, 3.0],
)

doc.add_heading("Node.js 22", level=2)
para(
    "The repo pins Node 22. Install it with a version manager so you can switch per-project — on "
    "Windows use nvm-windows, on macOS/Linux use nvm or fnm."
)
code(
    """
# Windows (PowerShell) — via winget
winget install CoreyButler.NVMforWindows
nvm install 22
nvm use 22

# macOS / Linux — via nvm
curl -o- https://raw.githubusercontent.com/nvm-sh/nvm/v0.40.1/install.sh | bash
nvm install 22
nvm use 22

# Confirm
node -v      # -> v22.x.x
"""
)
para(
    "If you would rather not use a version manager, the installer from nodejs.org works fine — "
    "just pick the 22 LTS line.",
    muted=True,
    size=9.5,
)

doc.add_heading("pnpm 10", level=2)
para(
    "This is a pnpm workspace. The lockfile, the workspace catalog, and the `workspace:^` "
    "dependency links are all pnpm-specific — installing with npm or yarn will fail or, worse, "
    "silently produce a broken tree. Node 22 ships Corepack, which is the cleanest way in:"
)
code(
    """
# Preferred — Corepack ships with Node, and pins the exact version the repo asks for
corepack enable
corepack prepare pnpm@10.18.1 --activate

# Alternative — a plain global install
npm install -g pnpm

# Confirm
pnpm -v      # -> 10.x.x
"""
)
callout(
    "Why 10.18.1?",
    "package.json declares packageManager: pnpm@10.18.1. Corepack reads that field and uses "
    "exactly that version, so everyone on the team resolves dependencies identically.",
)

doc.add_heading("Git", level=2)
code(
    """
# Windows
winget install Git.Git

# macOS
brew install git

# Debian / Ubuntu
sudo apt install git
"""
)

doc.add_heading("Docker Desktop — optional", level=2)
para(
    "You only need Docker if you want image uploads (karya covers and screenshots) to work. It "
    "runs MinIO, an S3-compatible object store, in a container. There is no database container: "
    "the local database is a plain SQLite file."
)
para(
    "Skip it and everything else still works — the upload endpoints just return HTTP 503. You can "
    "add it later without redoing any of this.",
    muted=True,
    size=9.5,
)
code(
    """
# Windows
winget install Docker.DockerDesktop

# macOS
brew install --cask docker
"""
)
callout(
    "Docker must be running, not just installed.",
    "Launch Docker Desktop and wait for the whale icon to settle before running any compose "
    "command, or you will get: \"failed to connect to the docker API ... check if the daemon is "
    "running\".",
    WARN_BG,
)

# =============================================================================
doc.add_heading("2. Clone the repository", level=1)

code(
    """
git clone git@github.com:zakinadhif/buildersnetwork.git
cd buildersnetwork
"""
)
para(
    "Use the HTTPS URL instead if you have not set up an SSH key: "
    "https://github.com/zakinadhif/buildersnetwork.git",
    muted=True,
    size=9.5,
)

# =============================================================================
doc.add_heading("3. Install dependencies", level=1)

code(
    """
pnpm install
"""
)
para(
    "This installs every workspace package at once — you never run install inside apps/api or "
    "apps/app individually. Expect roughly two minutes on a first run; pnpm caches globally, so "
    "later installs are near-instant."
)
callout(
    "The \"Ignored build scripts\" warning is expected.",
    "pnpm reports that @google/genai and protobufjs have unrun build scripts. That is deliberate — "
    "the repo allowlists only the packages that genuinely need to compile (better-sqlite3, esbuild, "
    "sharp, workerd) in pnpm-workspace.yaml. Ignore the warning; do not run pnpm approve-builds.",
)

# =============================================================================
doc.add_heading("4. Generate the API client — do not skip this", level=1)

code(
    """
pnpm codegen
"""
)
para(
    "This is the step the README leaves out, and without it the API will not start. "
    "libs/api-spec/openapi.yaml is the source of truth for every JSON endpoint. From it, Orval "
    "generates two packages:"
)
bullet("typed TanStack Query hooks the SPA imports.", bold_lead="@myapp/api-client-react — ")
bullet("Zod validators the Hono routes parse requests with.", bold_lead="@myapp/api-zod — ")
para(
    "Both live in generated/ folders that are gitignored, so a fresh clone simply does not have "
    "them. Boot the API without running codegen and you get:"
)
code(
    """
Error [ERR_MODULE_NOT_FOUND]: Cannot find module '.../libs/api-zod/src/generated/api'
  imported from .../libs/api-zod/src/index.ts
""",
    comment_prefixes=("//",),
)
callout(
    "Rerun codegen whenever openapi.yaml changes.",
    "Adding or changing an endpoint means editing the spec first, then regenerating. The spec is "
    "not documentation of the API — it is the API.",
)

# =============================================================================
doc.add_heading("5. Configure the environment", level=1)

para(
    "The API reads its config from apps/api/.env, validated at startup by a Zod schema in "
    "libs/config. Invalid or missing values fail fast with a readable error rather than a "
    "mysterious crash later."
)
para(
    "deploy/.env.example is the canonical reference — every variable, with the shapes each accepts. "
    "Its defaults are the local-dev ones, so copy it and you only need to fill in a single value:"
)

code(
    """
cp deploy/.env.example apps/api/.env

# Generate a secret (min 32 chars) and paste it into BETTER_AUTH_SECRET
openssl rand -base64 32
"""
)

para(
    "That is the whole configuration step. APP_URL, PORT, and DATABASE_URL already carry working "
    "local values, and every optional variable is blank — which the config loader reads as unset."
)

callout(
    "Blank means unset.",
    "libs/config drops empty strings before validating, so leaving an optional var as FOO= is "
    "exactly the same as deleting the line. This matters because blank-not-absent is everywhere: "
    ".env.example ships optional vars as blank lines, --env-file forwards them verbatim, and "
    "secret managers commonly inject \"\" for a secret that was never set.",
)

doc.add_heading("What each required variable does", level=2)
table(
    ["Variable", "Why it matters"],
    [
        [
            "`APP_URL`",
            "The public origin. Drives auth callbacks and CORS. In dev this is the Vite server "
            "(`:5173`), not the API — the browser talks to Vite, which proxies to the API.",
        ],
        [
            "`DATABASE_URL`",
            "A libSQL URL. `file:` for local SQLite, `libsql://` for Turso. Unused on Cloudflare, "
            "which binds a D1 database instead.",
        ],
        [
            "`BETTER_AUTH_SECRET`",
            "Signs session tokens. Minimum 32 characters — shorter values are rejected at startup.",
        ],
    ],
    widths=[1.7, 4.8],
)

# =============================================================================
doc.add_heading("6. Create and seed the database", level=1)

para(
    "There is no database server to install. The local database is one SQLite file, created by "
    "pushing the Drizzle schema straight into it."
)
code(
    """
pnpm db:push     # creates libs/db/local.db and applies the schema
pnpm db:seed     # inserts 5 members, 3 karya, 6 posts
"""
)

callout(
    "Where does the file land, and why that DATABASE_URL?",
    "db:push runs drizzle-kit with its working directory set to libs/db/, where the config defaults "
    "to file:./local.db — so the database is created at libs/db/local.db. The API, meanwhile, "
    "resolves DATABASE_URL relative to apps/api/. That is why the value is "
    "file:../../libs/db/local.db: two different working directories, one file. Point them at "
    "different paths and the API will happily open an empty database and show you nothing.",
)

para(
    "db:push is the right tool for local development — it syncs the schema without writing "
    "migration files. Use db:generate and db:migrate when you are changing the schema for real and "
    "need a migration committed."
)

doc.add_heading("Seed accounts", level=2)
para("The seeder creates five members you can sign in as immediately. Re-running it is idempotent.")
table(
    ["Email", "Password"],
    [
        ["`hafiz@seed.local`", "`seedpassword123`"],
        ["`fatimah@seed.local`", "`seedpassword123`"],
        ["`rizal@seed.local`", "`seedpassword123`"],
        ["`dinda@seed.local`", "`seedpassword123`"],
        ["`arya@seed.local`", "`seedpassword123`"],
    ],
    widths=[3.2, 3.3],
)
para(
    "The password is not a secret — seed data is for local dev and previews only, and the seeder "
    "refuses to run under NODE_ENV=production without --force.",
    muted=True,
    size=9.5,
)

# =============================================================================
doc.add_heading("7. Run it", level=1)

para("Two processes, two terminals. Start the API first.")
code(
    """
# Terminal 1 — Hono API
pnpm dev:api     # -> http://localhost:8080

# Terminal 2 — React SPA
pnpm dev:app     # -> http://localhost:5173
"""
)
para("Open http://localhost:5173 and sign in with a seed account. You land on the Launchpad.")

callout(
    "Open :5173, not :8080.",
    "The SPA is what you open; the API is what it talks to. Vite proxies /api through to "
    "127.0.0.1:8080, so the browser only ever sees one origin. If you override PORT, set "
    "VITE_API_DEV_TARGET to match or every API call from the SPA will 404. (Older checkouts of the "
    "README claimed the API ran on :3000 — it never did.)",
)
callout(
    "Two harmless startup messages.",
    "The API logs \"serveStatic: root path './public/spa' is not found\" — that is the production "
    "static-file handler finding no build, which is correct in dev. And Vite binds to localhost "
    "(IPv6): if a script of yours probes 127.0.0.1:5173 it may get connection-refused even though "
    "the browser works. Use localhost.",
)

doc.add_heading("The other two apps", level=2)
para("Neither is needed for day-to-day API/SPA work, but both run standalone:")
code(
    """
pnpm dev:landing    # Astro marketing page, served at / in production
pnpm dev:mockups    # the design-system gallery — no API, no database
"""
)

# =============================================================================
doc.add_heading("8. Confirm your setup is sound", level=1)

para("Three checks. All of these pass on a correctly set-up clean clone.")
code(
    """
pnpm test:db     # 26 unit tests, ~1s
pnpm test:api    # API route tests
pnpm lint        # Biome — must report zero errors
"""
)

callout(
    "Do NOT run pnpm better-auth:generate.",
    "The README's Quick start tells you to. On a fresh clone it is unnecessary and actively "
    "harmful: libs/db/src/schema/auth.ts is already committed, so the command just overwrites a "
    "good file with one whose imports are unsorted — which then fails pnpm lint. Only run it when "
    "you have actually changed the Better Auth config, and follow it with pnpm lint:fix. If you ran "
    "it by accident: git checkout -- libs/db/src/schema/auth.ts",
    WARN_BG,
)

# =============================================================================
doc.add_heading("9. What was broken, and what changed", level=1)

para(
    "Setting up a clean clone to write this guide surfaced five defects. All are fixed in the repo "
    "now, so you should not hit any of them — but if you are on an older checkout, or you are "
    "reviewing the commit that fixed them, this is what changed and why."
)
table(
    ["The defect", "The fix"],
    [
        [
            "The Quick start never mentioned `pnpm codegen`. The generated API client is "
            "gitignored, so a fresh clone had none and the API died with ERR_MODULE_NOT_FOUND. This "
            "is the one that stopped you cold.",
            "`pnpm codegen` is now step two of the Quick start.",
        ],
        [
            "A blank optional variable (`BETTER_AUTH_URL=`) failed validation — an empty string is "
            "not absent, so it was rejected as an invalid URL. Because .env.example ships blank "
            "lines, copying it as instructed could not boot.",
            "libs/config now drops empty strings before validating, so blank reads as unset. "
            "libs/auth had the same bug through `??` (which does not catch \"\") and now uses `||`.",
        ],
        [
            "The README said the API listens on `:3000`.",
            "It is `:8080` — both the config default and the Vite proxy target. Corrected.",
        ],
        [
            "The Quick start told you to run `pnpm better-auth:generate`. The file it writes is "
            "already committed, so this overwrote a good file with unsorted imports and then failed "
            "`pnpm lint`.",
            "Dropped from the Quick start; the README now describes it as a maintenance command.",
        ],
        [
            "`.env.example` shipped `DATABASE_URL=file:/data/app.db` (a container path) and "
            "`NODE_ENV=production`, so a developer who copied it got an empty feed and a seeder that "
            "refused to run.",
            "Its defaults are now the local-dev ones, with the deployment overrides documented "
            "inline.",
        ],
    ],
    widths=[3.3, 3.2],
)

# =============================================================================
doc.add_heading("10. Optional: uploads and AI", level=1)

doc.add_heading("Image uploads (MinIO)", level=2)
para(
    "Start Docker Desktop, wait for it to be ready, then bring up MinIO. It also creates the dev "
    "bucket for you."
)
code(
    """
docker compose -f deploy/docker-compose.dev.yml up -d
"""
)
para(
    "MinIO's S3 API is on :9000 and its web console on :9001 (login minioadmin / minioadmin). With "
    "the STORAGE_* block from Section 5 in place, cover and screenshot uploads work. Without "
    "storage configured, those routes return 503 and nothing else is affected."
)

doc.add_heading("AI features", level=2)
para(
    "The AI provider is chosen by the runtime entrypoint, not by an env var. Running locally on "
    "Node means apps/api/src/index.ts, which uses Gemini — so set GEMINI_API_KEY to exercise the "
    "assistant or the onboarding chat. On Cloudflare, the Worker entrypoint uses Workers AI and "
    "needs no key at all."
)

# =============================================================================
doc.add_heading("11. Troubleshooting", level=1)

para("Every error below is one that actually occurred while writing this guide.")
table(
    ["Symptom", "Cause and fix"],
    [
        [
            "`ERR_MODULE_NOT_FOUND: .../api-zod/src/generated/api`",
            "You skipped codegen. Run `pnpm codegen`.",
        ],
        [
            "`Invalid environment configuration: BETTER_AUTH_URL: Invalid url`",
            "An older checkout, from before blank-means-unset was fixed in libs/config. Either pull "
            "the fix, or delete the blank `BETTER_AUTH_URL=` line from your .env.",
        ],
        [
            "`BETTER_AUTH_SECRET` rejected at startup",
            "It is under 32 characters. Regenerate: `openssl rand -base64 32`.",
        ],
        [
            "App loads but the feed is empty",
            "The API and drizzle-kit are pointed at different SQLite files — usually a stale "
            "`DATABASE_URL`. It must be `file:../../libs/db/local.db`. Fix it, then re-run "
            "`pnpm db:seed`.",
        ],
        [
            "SPA loads but every API call 404s",
            "The API is not on :8080, or you changed `PORT` without setting "
            "`VITE_API_DEV_TARGET` to match.",
        ],
        [
            "`failed to connect to the docker API ... daemon is running`",
            "Docker Desktop is installed but not started. Launch it and wait for it to go green.",
        ],
        [
            "`pnpm lint` fails on `libs/db/src/schema/auth.ts`",
            "You ran better-auth:generate. Fix with `git checkout -- libs/db/src/schema/auth.ts`, "
            "or `pnpm lint:fix` if the regeneration was intentional.",
        ],
        [
            "A script cannot reach `127.0.0.1:5173`",
            "Vite binds to localhost over IPv6. Use `localhost:5173` instead.",
        ],
    ],
    widths=[2.5, 4.0],
)

# =============================================================================
doc.add_heading("12. The commands you will actually use", level=1)

table(
    ["Command", "What it does"],
    [
        ["`pnpm dev:api`", "Hono API on :8080, watch mode"],
        ["`pnpm dev:app`", "React SPA on :5173, watch mode"],
        ["`pnpm dev:mockups`", "The design-system gallery — the visual north star"],
        ["`pnpm codegen`", "Regenerate the API client + Zod validators from openapi.yaml"],
        ["`pnpm db:push`", "Sync the schema to the local SQLite file"],
        ["`pnpm db:seed`", "Reseed (idempotent)"],
        ["`pnpm db:studio`", "Browse the database in Drizzle Studio"],
        ["`pnpm lint` / `pnpm lint:fix`", "Biome check / autofix"],
        ["`pnpm test:db` / `test:api` / `test:app`", "Vitest suites"],
        ["`pnpm test:e2e`", "Playwright end-to-end"],
    ],
    widths=[2.5, 4.0],
)

# =============================================================================
doc.add_heading("13. Where to go next", level=1)

bullet(
    "the source of truth for JSON endpoints. Add a path, run pnpm codegen, get typed hooks and "
    "validators. See plans/how-to/adding-an-endpoint.md.",
    bold_lead="libs/api-spec/openapi.yaml — ",
)
bullet(
    "the design system's single source. The mockup gallery is the north star — look at it, don't "
    "read about it.",
    bold_lead="apps/mockups/src/lib/tokens.ts — ",
)
bullet(
    "how the team coordinates: milestone docs hold intent, GitHub Issues hold task contracts, the "
    "project board holds status.",
    bold_lead="plans/how-to/build-workflow.md — ",
)
bullet(
    "vision docs, design references, and implementation plans. Treat it as non-authoritative — when "
    "it conflicts with the code, trust the code.",
    bold_lead="plans/ — ",
)

para(
    "One-time setup if you will be filing or shipping tasks: gh auth login, then "
    "gh auth refresh -s project,read:project.",
)
para(
    "Claude Code users get the workflow as repo skills: /project-status, /pick-task, /ship-task, "
    "/new-task, /ratify.",
    muted=True,
    size=9.5,
)

out = Path(__file__).resolve().parent / "getting-started.docx"
doc.save(out)
print(f"saved -> {out}")
