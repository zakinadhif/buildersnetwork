"""Membangun panduan memulai Al-Fath Berkarya (Bahasa Indonesia) sebagai .docx.

File .docx-nya adalah artefak build — edit script ini, bukan file Word-nya, atau
perubahanmu hilang saat script dijalankan lagi. Edisi bahasa Inggrisnya ada di
build-guide-en.py; keduanya dijaga sejalan, jadi perubahan di sini biasanya perlu
diterapkan juga di sana.

    pip install python-docx
    python docs/build-guide-id.py    # -> docs/panduan-memulai.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x7A, 0x2E, 0x2E)      # maroon, dari brand Al-Fath
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
CODE_BG = "F2F1EF"
NOTE_BG = "FBF3E4"
WARN_BG = "FBE9E7"

doc = Document()

# --- base styles -------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for name, size in (("Heading 1", 18), ("Heading 2", 13.5), ("Heading 3", 11.5)):
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = ACCENT
    st.paragraph_format.space_before = Pt(14)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True


def shade(paragraph, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    paragraph._p.get_or_add_pPr().append(shd)


def code(lines, comment_prefixes=("#",)):
    if isinstance(lines, str):
        lines = lines.strip("\n").split("\n")
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.right_indent = Inches(0.10)
    pf.space_before = Pt(6)
    pf.space_after = Pt(9)
    pf.line_spacing = 1.05
    shade(p, CODE_BG)
    for i, line in enumerate(lines):
        run = p.add_run(("\n" if i else "") + line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        is_comment = line.lstrip().startswith(comment_prefixes)
        run.font.color.rgb = MUTED if is_comment else RGBColor(0x1A, 0x1A, 0x1A)
    return p


def callout(label, text, bg=NOTE_BG):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.right_indent = Inches(0.10)
    pf.space_before = Pt(6)
    pf.space_after = Pt(9)
    shade(p, bg)
    r = p.add_run(f"{label}  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = ACCENT if bg is NOTE_BG else RGBColor(0xA3, 0x2A, 0x1C)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p


def para(text, bold=False, italic=False, muted=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if muted:
        r.font.color.rgb = MUTED
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            for j, chunk in enumerate(val.split("`")):
                if not chunk:
                    continue
                run = p.add_run(chunk)
                if j % 2:
                    run.font.name = "Consolas"
                    run.font.size = Pt(8.5)
                else:
                    run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# =============================================================================
# Judul
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("Al-Fath Berkarya")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = ACCENT
r.font.name = "Calibri"

sub = doc.add_paragraph()
r = sub.add_run("Panduan Memulai — menjalankan project dari nol untuk developer")
r.font.size = Pt(13)
r.font.color.rgb = MUTED

meta = doc.add_paragraph()
r = meta.add_run(
    "Diverifikasi end-to-end di Windows 11 pada commit c9b3b5f — 14 Juli 2026.\n"
    "Semua perintah di panduan ini benar-benar dijalankan pada clone yang bersih, dan setiap error "
    "di bagian Troubleshooting adalah error yang sungguh-sungguh muncul saat penyusunan panduan ini."
)
r.font.size = Pt(9.5)
r.italic = True
r.font.color.rgb = MUTED

doc.add_paragraph()

para(
    "Al-Fath Berkarya adalah platform komunitas untuk mahasiswa builder di Telkom University. "
    "Bentuknya monorepo pnpm: sebuah API Hono, SPA React 19, landing page Astro, dan galeri mockup "
    "statis, yang berbagi sekumpulan library (database, auth, AI, storage, config). Codebase yang "
    "sama bisa dideploy ke Cloudflare Workers maupun ke Node/Docker — yang menentukan provider AI "
    "dan client database adalah entrypoint runtime-nya, bukan environment variable."
)
para(
    "Panduan ini membawa kamu dari belum meng-install apa pun sampai aplikasi berjalan dengan data "
    "seed yang bisa langsung kamu login. Siapkan sekitar 20 menit, sebagian besar cuma menunggu "
    "proses download."
)

callout(
    "Sedikit catatan sejarah.",
    "Penyusunan panduan ini dilakukan dengan benar-benar men-setup clone yang bersih, dan dari situ "
    "muncul lima cacat pada alur yang didokumentasikan — termasuk satu yang membuat clone baru sama "
    "sekali tidak bisa berjalan. Semuanya sudah diperbaiki di repo, dan Quick start di README kini "
    "sudah benar per commit c9b3b5f. Panduan ini membahas hal yang sama secara lebih mendalam. "
    "Bagian 9 mencatat apa saja yang dulu keliru, kalau-kalau kamu memakai checkout yang lebih lama.",
)

# =============================================================================
doc.add_heading("1. Yang perlu di-install", level=1)

para(
    "Ada empat tools. Node dan pnpm wajib; Git adalah cara kamu mengambil kodenya; Docker sifatnya "
    "opsional dan hanya dibutuhkan untuk fitur upload file."
)

table(
    ["Tools", "Versi", "Wajib?"],
    [
        ["Node.js", "22 atau lebih baru", "Ya"],
        ["pnpm", "10 atau lebih baru", "Ya — npm dan yarn tidak akan jalan"],
        ["Git", "versi terbaru mana pun", "Ya"],
        ["Docker Desktop", "versi terbaru mana pun", "Opsional — hanya untuk upload gambar (MinIO)"],
    ],
    widths=[1.6, 1.9, 3.0],
)

doc.add_heading("Node.js 22", level=2)
para(
    "Repo ini dipatok di Node 22. Install pakai version manager supaya kamu bisa berganti versi "
    "per-project — di Windows pakai nvm-windows, di macOS/Linux pakai nvm atau fnm."
)
code(
    """
# Windows (PowerShell) — lewat winget
winget install CoreyButler.NVMforWindows
nvm install 22
nvm use 22

# macOS / Linux — lewat nvm
curl -o- https://raw.githubusercontent.com/nvm-sh/nvm/v0.40.1/install.sh | bash
nvm install 22
nvm use 22

# Cek hasilnya
node -v      # -> v22.x.x
"""
)
para(
    "Kalau kamu tidak mau pakai version manager, installer dari nodejs.org juga tidak masalah — "
    "pastikan saja memilih jalur 22 LTS.",
    muted=True,
    size=9.5,
)

doc.add_heading("pnpm 10", level=2)
para(
    "Project ini adalah pnpm workspace. Lockfile, workspace catalog, dan dependency link `workspace:^` "
    "semuanya khas pnpm — meng-install dengan npm atau yarn akan gagal, atau lebih buruk lagi, "
    "diam-diam menghasilkan dependency tree yang rusak. Node 22 sudah membawa Corepack, dan itu cara "
    "paling bersih:"
)
code(
    """
# Disarankan — Corepack bawaan Node, dan otomatis memakai versi persis yang diminta repo
corepack enable
corepack prepare pnpm@10.18.1 --activate

# Alternatif — install global biasa
npm install -g pnpm

# Cek hasilnya
pnpm -v      # -> 10.x.x
"""
)
callout(
    "Kenapa harus 10.18.1?",
    "Di package.json tertulis packageManager: pnpm@10.18.1. Corepack membaca field itu dan memakai "
    "versi tersebut persis, sehingga semua orang di tim me-resolve dependency dengan cara yang sama.",
)

doc.add_heading("Git", level=2)
code(
    """
# Windows
winget install Git.Git

# macOS
brew install git

# Debian / Ubuntu
sudo apt install git
"""
)

doc.add_heading("Docker Desktop — opsional", level=2)
para(
    "Docker hanya kamu butuhkan kalau ingin fitur upload gambar (cover dan screenshot karya) "
    "berfungsi. Docker menjalankan MinIO, object storage yang kompatibel dengan S3, di dalam "
    "container. Tidak ada container database: database lokal cuma sebuah file SQLite biasa."
)
para(
    "Kalau dilewati, semua bagian lain tetap jalan — endpoint upload-nya saja yang mengembalikan "
    "HTTP 503. Kamu bisa menambahkannya nanti tanpa perlu mengulang langkah mana pun.",
    muted=True,
    size=9.5,
)
code(
    """
# Windows
winget install Docker.DockerDesktop

# macOS
brew install --cask docker
"""
)
callout(
    "Docker harus benar-benar berjalan, bukan sekadar ter-install.",
    "Buka Docker Desktop dan tunggu sampai ikon paus-nya tenang sebelum menjalankan perintah compose "
    "apa pun. Kalau tidak, kamu akan dapat: \"failed to connect to the docker API ... check if the "
    "daemon is running\".",
    WARN_BG,
)

# =============================================================================
doc.add_heading("2. Clone repository", level=1)

code(
    """
git clone git@github.com:zakinadhif/buildersnetwork.git
cd buildersnetwork
"""
)
para(
    "Pakai URL HTTPS kalau kamu belum menyiapkan SSH key: "
    "https://github.com/zakinadhif/buildersnetwork.git",
    muted=True,
    size=9.5,
)

# =============================================================================
doc.add_heading("3. Install dependency", level=1)

code(
    """
pnpm install
"""
)
para(
    "Perintah ini meng-install seluruh package di workspace sekaligus — kamu tidak perlu menjalankan "
    "install di dalam apps/api atau apps/app satu per satu. Perkirakan sekitar dua menit untuk "
    "install pertama; pnpm menyimpan cache secara global, jadi install berikutnya jauh lebih cepat."
)
callout(
    "Peringatan \"Ignored build scripts\" itu normal.",
    "pnpm akan melaporkan bahwa @google/genai dan protobufjs punya build script yang tidak "
    "dijalankan. Itu memang disengaja — repo hanya mengizinkan package yang benar-benar perlu "
    "dikompilasi (better-sqlite3, esbuild, sharp, workerd) lewat pnpm-workspace.yaml. Abaikan "
    "peringatannya; jangan menjalankan pnpm approve-builds.",
)

# =============================================================================
doc.add_heading("4. Generate API client — jangan dilewati", level=1)

code(
    """
pnpm codegen
"""
)
para(
    "Inilah langkah yang tidak disebutkan README, dan tanpa ini API tidak akan bisa start. "
    "libs/api-spec/openapi.yaml adalah sumber kebenaran untuk semua endpoint JSON. Dari file itu, "
    "Orval meng-generate dua package:"
)
bullet("hooks TanStack Query bertipe, yang di-import oleh SPA.", bold_lead="@myapp/api-client-react — ")
bullet("validator Zod, yang dipakai route Hono untuk mem-parse request.", bold_lead="@myapp/api-zod — ")
para(
    "Keduanya berada di folder generated/ yang masuk gitignore, jadi clone yang baru memang belum "
    "memilikinya. Kalau kamu menjalankan API tanpa codegen, hasilnya:"
)
code(
    """
Error [ERR_MODULE_NOT_FOUND]: Cannot find module '.../libs/api-zod/src/generated/api'
  imported from .../libs/api-zod/src/index.ts
""",
    comment_prefixes=("//",),
)
callout(
    "Jalankan codegen lagi setiap kali openapi.yaml berubah.",
    "Menambah atau mengubah endpoint berarti mengedit spec-nya dulu, baru generate ulang. Spec ini "
    "bukan dokumentasi tentang API — spec ini adalah API-nya.",
)

# =============================================================================
doc.add_heading("5. Menyiapkan environment", level=1)

para(
    "API membaca konfigurasinya dari apps/api/.env, yang divalidasi saat startup oleh schema Zod di "
    "libs/config. Nilai yang salah atau kurang akan langsung gagal dengan pesan error yang jelas, "
    "bukan crash misterius di kemudian hari."
)
para(
    "deploy/.env.example adalah referensi resminya — memuat semua variabel beserta bentuk nilai yang "
    "diterima masing-masing. Nilai default-nya sudah disetel untuk development lokal, jadi cukup "
    "salin file itu dan kamu hanya perlu mengisi satu nilai:"
)

code(
    """
cp deploy/.env.example apps/api/.env

# Buat secret (minimal 32 karakter), lalu tempel ke BETTER_AUTH_SECRET
openssl rand -base64 32
"""
)

para(
    "Hanya itu langkah konfigurasinya. APP_URL, PORT, dan DATABASE_URL sudah berisi nilai lokal yang "
    "berfungsi, dan semua variabel opsional dibiarkan kosong — yang oleh config loader dibaca "
    "sebagai tidak-diset."
)

callout(
    "Kosong berarti tidak diset.",
    "libs/config membuang string kosong sebelum melakukan validasi, jadi membiarkan variabel opsional "
    "sebagai FOO= sama persis dengan menghapus barisnya. Ini penting karena kondisi kosong-tapi-ada "
    "itu umum terjadi: .env.example memang mengirim variabel opsional sebagai baris kosong, "
    "--env-file meneruskannya apa adanya, dan secret manager sering menyuntikkan \"\" untuk secret "
    "yang belum pernah diisi.",
)

doc.add_heading("Fungsi tiap variabel wajib", level=2)
table(
    ["Variabel", "Kenapa penting"],
    [
        [
            "`APP_URL`",
            "Origin publik aplikasi. Dipakai untuk auth callback dan CORS. Saat development ini "
            "adalah server Vite (`:5173`), bukan API — browser berbicara ke Vite, lalu Vite "
            "mem-proxy ke API.",
        ],
        [
            "`DATABASE_URL`",
            "Sebuah URL libSQL. `file:` untuk SQLite lokal, `libsql://` untuk Turso. Tidak dipakai "
            "di Cloudflare, yang memakai binding database D1.",
        ],
        [
            "`BETTER_AUTH_SECRET`",
            "Menandatangani session token. Minimal 32 karakter — nilai yang lebih pendek ditolak "
            "saat startup.",
        ],
    ],
    widths=[1.7, 4.8],
)

# =============================================================================
doc.add_heading("6. Membuat dan mengisi database", level=1)

para(
    "Tidak ada database server yang perlu di-install. Database lokal hanyalah satu file SQLite, yang "
    "dibuat dengan mendorong (push) schema Drizzle langsung ke dalamnya."
)
code(
    """
pnpm db:push     # membuat libs/db/local.db dan menerapkan schema
pnpm db:seed     # mengisi 5 member, 3 karya, 6 post
"""
)

callout(
    "Di mana file-nya dibuat, dan kenapa DATABASE_URL-nya begitu?",
    "db:push menjalankan drizzle-kit dengan working directory libs/db/, yang konfigurasinya "
    "default ke file:./local.db — jadi database dibuat di libs/db/local.db. Sementara itu, API "
    "me-resolve DATABASE_URL relatif terhadap apps/api/. Itulah sebabnya nilainya "
    "file:../../libs/db/local.db: dua working directory berbeda, satu file yang sama. Kalau keduanya "
    "menunjuk path berbeda, API akan tenang-tenang saja membuka database kosong dan tidak menampilkan "
    "apa pun.",
)

para(
    "db:push adalah pilihan yang tepat untuk development lokal — ia menyinkronkan schema tanpa "
    "menulis file migration. Gunakan db:generate dan db:migrate kalau kamu benar-benar mengubah "
    "schema dan perlu migration yang di-commit."
)

doc.add_heading("Akun seed", level=2)
para(
    "Seeder membuat lima member yang bisa langsung kamu pakai login. Menjalankannya berulang kali "
    "aman (idempotent)."
)
table(
    ["Email", "Password"],
    [
        ["`hafiz@seed.local`", "`seedpassword123`"],
        ["`fatimah@seed.local`", "`seedpassword123`"],
        ["`rizal@seed.local`", "`seedpassword123`"],
        ["`dinda@seed.local`", "`seedpassword123`"],
        ["`arya@seed.local`", "`seedpassword123`"],
    ],
    widths=[3.2, 3.3],
)
para(
    "Password ini bukan rahasia — data seed hanya untuk development lokal dan preview, dan seeder-nya "
    "menolak berjalan di NODE_ENV=production tanpa flag --force.",
    muted=True,
    size=9.5,
)

# =============================================================================
doc.add_heading("7. Menjalankan aplikasi", level=1)

para("Dua proses, dua terminal. Jalankan API lebih dulu.")
code(
    """
# Terminal 1 — API Hono
pnpm dev:api     # -> http://localhost:8080

# Terminal 2 — SPA React
pnpm dev:app     # -> http://localhost:5173
"""
)
para(
    "Buka http://localhost:5173 lalu login dengan salah satu akun seed. Kamu akan mendarat di "
    "Launchpad."
)

callout(
    "Yang kamu buka adalah :5173, bukan :8080.",
    "SPA-lah yang kamu buka; API adalah lawan bicaranya. Vite mem-proxy /api ke 127.0.0.1:8080, jadi "
    "browser hanya pernah melihat satu origin. Kalau kamu mengubah PORT, setel juga "
    "VITE_API_DEV_TARGET agar cocok — kalau tidak, semua panggilan API dari SPA akan 404. (Checkout "
    "README yang lama menyebut API berjalan di :3000 — padahal tidak pernah begitu.)",
)
callout(
    "Dua pesan startup yang tidak perlu dikhawatirkan.",
    "API akan mencatat \"serveStatic: root path './public/spa' is not found\" — itu handler file "
    "statis untuk production yang tidak menemukan hasil build, dan itu wajar saat development. Lalu, "
    "Vite bind ke localhost (IPv6): kalau ada script kamu yang mengecek 127.0.0.1:5173, bisa saja "
    "connection-refused padahal di browser normal. Gunakan localhost.",
)

doc.add_heading("Dua aplikasi lainnya", level=2)
para(
    "Keduanya tidak diperlukan untuk kerja harian di API/SPA, tapi bisa dijalankan sendiri:"
)
code(
    """
pnpm dev:landing    # landing page Astro, disajikan di / saat production
pnpm dev:mockups    # galeri design system — tanpa API, tanpa database
"""
)

# =============================================================================
doc.add_heading("8. Memastikan setup kamu benar", level=1)

para("Tiga pemeriksaan. Semuanya lolos pada clone bersih yang di-setup dengan benar.")
code(
    """
pnpm test:db     # 26 unit test, ~1 detik
pnpm test:api    # test untuk route API
pnpm lint        # Biome — harus nol error
"""
)

callout(
    "JANGAN jalankan pnpm better-auth:generate.",
    "README menyuruh kamu menjalankannya. Pada clone yang baru, perintah itu tidak perlu dan justru "
    "merugikan: libs/db/src/schema/auth.ts sudah ikut di-commit, jadi perintah tersebut hanya menimpa "
    "file yang sudah benar dengan versi yang urutan import-nya berantakan — dan itu membuat pnpm lint "
    "gagal. Jalankan hanya kalau kamu memang mengubah konfigurasi Better Auth, lalu susul dengan "
    "pnpm lint:fix. Kalau terlanjur menjalankannya: git checkout -- libs/db/src/schema/auth.ts",
    WARN_BG,
)

# =============================================================================
doc.add_heading("9. Apa yang dulu rusak, dan apa yang berubah", level=1)

para(
    "Proses men-setup clone bersih untuk menyusun panduan ini memunculkan lima cacat. Semuanya sudah "
    "diperbaiki di repo, jadi seharusnya kamu tidak akan mengalaminya — tapi kalau kamu memakai "
    "checkout lama, atau sedang me-review commit yang memperbaikinya, inilah rinciannya."
)
table(
    ["Cacatnya", "Perbaikannya"],
    [
        [
            "Quick start tidak pernah menyebut `pnpm codegen`. Client hasil generate masuk gitignore, "
            "jadi clone baru tidak memilikinya dan API mati dengan ERR_MODULE_NOT_FOUND. Inilah yang "
            "membuat kamu mentok total.",
            "`pnpm codegen` kini menjadi langkah kedua di Quick start.",
        ],
        [
            "Variabel opsional yang kosong (`BETTER_AUTH_URL=`) gagal validasi — string kosong tidak "
            "sama dengan tidak-ada, sehingga ditolak sebagai URL yang tidak valid. Karena "
            ".env.example memang berisi baris kosong, menyalinnya sesuai instruksi justru tidak bisa "
            "boot.",
            "libs/config kini membuang string kosong sebelum validasi, jadi kosong dibaca sebagai "
            "tidak-diset. libs/auth punya bug yang sama lewat `??` (yang tidak menangkap \"\") dan "
            "kini memakai `||`.",
        ],
        [
            "README menyebut API berjalan di `:3000`.",
            "Yang benar `:8080` — baik sebagai default config maupun target proxy Vite. Sudah "
            "diperbaiki.",
        ],
        [
            "Quick start menyuruh menjalankan `pnpm better-auth:generate`. File yang ditulisnya sudah "
            "ikut di-commit, jadi perintah itu hanya menimpa file yang benar dengan versi yang urutan "
            "import-nya berantakan, lalu membuat `pnpm lint` gagal.",
            "Dihapus dari Quick start; README kini menjelaskannya sebagai perintah pemeliharaan.",
        ],
        [
            "`.env.example` membawa `DATABASE_URL=file:/data/app.db` (path di dalam container) dan "
            "`NODE_ENV=production`, sehingga developer yang menyalinnya mendapat feed kosong dan "
            "seeder yang menolak berjalan.",
            "Nilai default-nya kini disetel untuk development lokal, dengan override untuk deployment "
            "didokumentasikan langsung di file itu.",
        ],
    ],
    widths=[3.3, 3.2],
)

# =============================================================================
doc.add_heading("10. Opsional: upload dan AI", level=1)

doc.add_heading("Upload gambar (MinIO)", level=2)
para(
    "Jalankan Docker Desktop, tunggu sampai siap, lalu nyalakan MinIO. Perintah ini sekalian membuat "
    "bucket untuk development."
)
code(
    """
docker compose -f deploy/docker-compose.dev.yml up -d
"""
)
para(
    "S3 API milik MinIO ada di :9000 dan konsol web-nya di :9001 (login minioadmin / minioadmin). "
    "Dengan blok STORAGE_* dari Bagian 5 terpasang, upload cover dan screenshot akan berfungsi. Tanpa "
    "storage yang dikonfigurasi, route tersebut mengembalikan 503 dan tidak memengaruhi bagian lain."
)

doc.add_heading("Fitur AI", level=2)
para(
    "Provider AI ditentukan oleh entrypoint runtime, bukan oleh environment variable. Menjalankan "
    "aplikasi secara lokal di Node berarti memakai apps/api/src/index.ts, yang menggunakan Gemini — "
    "jadi isilah GEMINI_API_KEY kalau kamu ingin mencoba asisten atau onboarding chat. Di Cloudflare, "
    "entrypoint Worker memakai Workers AI dan sama sekali tidak butuh API key."
)

# =============================================================================
doc.add_heading("11. Troubleshooting", level=1)

para("Setiap error di bawah ini benar-benar muncul saat panduan ini disusun.")
table(
    ["Gejala", "Penyebab dan solusi"],
    [
        [
            "`ERR_MODULE_NOT_FOUND: .../api-zod/src/generated/api`",
            "Kamu melewatkan codegen. Jalankan `pnpm codegen`.",
        ],
        [
            "`Invalid environment configuration: BETTER_AUTH_URL: Invalid url`",
            "Checkout lama, dari sebelum perbaikan kosong-berarti-tidak-diset di libs/config. Pull "
            "perbaikannya, atau hapus baris `BETTER_AUTH_URL=` yang kosong dari .env kamu.",
        ],
        [
            "`BETTER_AUTH_SECRET` ditolak saat startup",
            "Panjangnya kurang dari 32 karakter. Buat ulang: `openssl rand -base64 32`.",
        ],
        [
            "Aplikasi terbuka tapi feed-nya kosong",
            "API dan drizzle-kit menunjuk file SQLite yang berbeda — biasanya karena `DATABASE_URL` "
            "yang usang. Nilainya harus `file:../../libs/db/local.db`. Perbaiki, lalu jalankan "
            "`pnpm db:seed` lagi.",
        ],
        [
            "SPA terbuka tapi semua panggilan API 404",
            "API tidak berjalan di :8080, atau kamu mengubah `PORT` tanpa menyetel "
            "`VITE_API_DEV_TARGET` agar cocok.",
        ],
        [
            "`failed to connect to the docker API ... daemon is running`",
            "Docker Desktop sudah ter-install tapi belum dijalankan. Buka aplikasinya dan tunggu "
            "sampai statusnya hijau.",
        ],
        [
            "`pnpm lint` gagal di `libs/db/src/schema/auth.ts`",
            "Kamu menjalankan better-auth:generate. Perbaiki dengan "
            "`git checkout -- libs/db/src/schema/auth.ts`, atau `pnpm lint:fix` kalau regenerasinya "
            "memang disengaja.",
        ],
        [
            "Script tidak bisa menjangkau `127.0.0.1:5173`",
            "Vite bind ke localhost lewat IPv6. Pakai `localhost:5173` saja.",
        ],
    ],
    widths=[2.5, 4.0],
)

# =============================================================================
doc.add_heading("12. Perintah yang akan sering kamu pakai", level=1)

table(
    ["Perintah", "Fungsinya"],
    [
        ["`pnpm dev:api`", "API Hono di :8080, mode watch"],
        ["`pnpm dev:app`", "SPA React di :5173, mode watch"],
        ["`pnpm dev:mockups`", "Galeri design system — acuan visual utama"],
        ["`pnpm codegen`", "Generate ulang API client + validator Zod dari openapi.yaml"],
        ["`pnpm db:push`", "Sinkronkan schema ke file SQLite lokal"],
        ["`pnpm db:seed`", "Isi ulang data seed (idempotent)"],
        ["`pnpm db:studio`", "Jelajahi database lewat Drizzle Studio"],
        ["`pnpm lint` / `pnpm lint:fix`", "Pemeriksaan / perbaikan otomatis Biome"],
        ["`pnpm test:db` / `test:api` / `test:app`", "Rangkaian test Vitest"],
        ["`pnpm test:e2e`", "Test end-to-end dengan Playwright"],
    ],
    widths=[2.5, 4.0],
)

# =============================================================================
doc.add_heading("13. Langkah selanjutnya", level=1)

bullet(
    "sumber kebenaran untuk semua endpoint JSON. Tambahkan path-nya, jalankan pnpm codegen, dan kamu "
    "dapat hooks bertipe beserta validatornya. Lihat plans/how-to/adding-an-endpoint.md.",
    bold_lead="libs/api-spec/openapi.yaml — ",
)
bullet(
    "satu-satunya sumber design system. Galeri mockup adalah acuan utamanya — lihat langsung, jangan "
    "cuma dibaca.",
    bold_lead="apps/mockups/src/lib/tokens.ts — ",
)
bullet(
    "cara tim berkoordinasi: dokumen milestone memuat maksud dan keputusan, GitHub Issues memuat "
    "kontrak tiap task, dan project board memuat statusnya.",
    bold_lead="plans/how-to/build-workflow.md — ",
)
bullet(
    "dokumen visi, referensi desain, dan rencana implementasi. Anggap isinya tidak otoritatif — kalau "
    "bertentangan dengan kode, percayai kodenya.",
    bold_lead="plans/ — ",
)

para(
    "Setup sekali di awal kalau kamu akan membuat atau mengirim task: gh auth login, lalu "
    "gh auth refresh -s project,read:project."
)
para(
    "Pengguna Claude Code mendapat alur kerja ini sebagai repo skills: /project-status, /pick-task, "
    "/ship-task, /new-task, /ratify.",
    muted=True,
    size=9.5,
)

out = Path(__file__).resolve().parent / "panduan-memulai.docx"
doc.save(out)
print(f"saved -> {out}")
