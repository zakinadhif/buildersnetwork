"""Shared docx styling used by the build-guide-*.py scripts.

One look for every guide: the Al-Fath maroon on headings, shaded monospace code
blocks, and two callout flavours (a note, and a warning). Import `new_doc()` to
get a styled Document plus the builders bound to it.

    pip install python-docx
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x7A, 0x2E, 0x2E)  # maroon, from the Al-Fath brand
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
WARN_INK = RGBColor(0xA3, 0x2A, 0x1C)

CODE_BG = "F2F1EF"
NOTE_BG = "FBF3E4"
WARN_BG = "FBE9E7"


def _shade(paragraph, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    paragraph._p.get_or_add_pPr().append(shd)


def new_doc():
    """A Document with the house styles applied, plus builders bound to it.

    Returns (doc, kit) where kit exposes .code/.callout/.para/.bullet/.table.
    """
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size in (("Heading 1", 18), ("Heading 2", 13.5), ("Heading 3", 11.5)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = ACCENT
        st.paragraph_format.space_before = Pt(14)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True

    return doc, _Kit(doc)


class _Kit:
    def __init__(self, doc):
        self.doc = doc

    def code(self, lines, comment_prefixes=("#",)):
        """A shaded monospace block. Comment lines render muted."""
        if isinstance(lines, str):
            lines = lines.strip("\n").split("\n")
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.18)
        pf.right_indent = Inches(0.10)
        pf.space_before = Pt(6)
        pf.space_after = Pt(9)
        pf.line_spacing = 1.05
        _shade(p, CODE_BG)
        for i, line in enumerate(lines):
            run = p.add_run(("\n" if i else "") + line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = (
                MUTED if line.lstrip().startswith(comment_prefixes) else INK
            )
        return p

    def callout(self, label, text, bg=NOTE_BG):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.18)
        pf.right_indent = Inches(0.10)
        pf.space_before = Pt(6)
        pf.space_after = Pt(9)
        _shade(p, bg)
        r = p.add_run(f"{label}  ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = ACCENT if bg == NOTE_BG else WARN_INK
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        return p

    def para(self, text, bold=False, italic=False, muted=False, size=10.5):
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if muted:
            r.font.color.rgb = MUTED
        return p

    def bullet(self, text, bold_lead=None):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if bold_lead:
            p.add_run(bold_lead).bold = True
        p.add_run(text)
        return p

    def table(self, headers, rows, widths=None):
        """A table. Backtick-wrapped spans inside a cell render as inline code."""
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(9.5)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                p = cells[i].paragraphs[0]
                for j, chunk in enumerate(val.split("`")):
                    if not chunk:
                        continue
                    run = p.add_run(chunk)
                    if j % 2:
                        run.font.name = "Consolas"
                        run.font.size = Pt(8.5)
                    else:
                        run.font.size = Pt(9.5)
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def title(self, name, subtitle, meta):
        p = self.doc.add_paragraph()
        r = p.add_run(name)
        r.font.size = Pt(28)
        r.bold = True
        r.font.color.rgb = ACCENT
        r.font.name = "Calibri"

        p = self.doc.add_paragraph()
        r = p.add_run(subtitle)
        r.font.size = Pt(13)
        r.font.color.rgb = MUTED

        p = self.doc.add_paragraph()
        r = p.add_run(meta)
        r.font.size = Pt(9.5)
        r.italic = True
        r.font.color.rgb = MUTED

        self.doc.add_paragraph()
